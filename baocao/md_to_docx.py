# -*- coding: utf-8 -*-
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = "noi_dung_bao_cao.md"
OUT = "bao_cao_TECHACADEMY.docx"

doc = Document()

# Base style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(13)

def add_inline(par, text):
    """Parse **bold**, `code` and emoji into runs."""
    # split keeping markers
    tokens = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(11)
        else:
            par.add_run(tok)

def is_table_sep(line):
    return bool(re.match(r"^\s*\|?[\s:\-\|]+\|?\s*$", line)) and "-" in line

lines = open(SRC, encoding="utf-8").read().split("\n")
i = 0
n = len(lines)
while i < n:
    line = lines[i].rstrip("\n")
    stripped = line.strip()

    # Horizontal rule
    if stripped == "---":
        i += 1
        continue

    # Code block
    if stripped.startswith("```"):
        i += 1
        code_lines = []
        while i < n and not lines[i].strip().startswith("```"):
            code_lines.append(lines[i]); i += 1
        i += 1  # skip closing ```
        for cl in code_lines:
            p = doc.add_paragraph()
            r = p.add_run(cl if cl else " ")
            r.font.name = "Consolas"; r.font.size = Pt(10.5)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(0)
        doc.add_paragraph()
        continue

    # Table
    if stripped.startswith("|") and i + 1 < n and is_table_sep(lines[i+1]):
        header = [c.strip() for c in stripped.strip("|").split("|")]
        i += 2  # skip header + separator
        rows = []
        while i < n and lines[i].strip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            rows.append(cells); i += 1
        tbl = doc.add_table(rows=1, cols=len(header))
        tbl.style = "Light Grid Accent 1"
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = tbl.rows[0].cells
        for j, h in enumerate(header):
            hdr[j].text = ""
            p = hdr[j].paragraphs[0]
            add_inline(p, h)
            for rr in p.runs:
                rr.bold = True
        for row in rows:
            rc = tbl.add_row().cells
            for j in range(len(header)):
                val = row[j] if j < len(row) else ""
                rc[j].text = ""
                add_inline(rc[j].paragraphs[0], val)
        doc.add_paragraph()
        continue

    # Headings
    m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
    if m:
        level = len(m.group(1))
        txt = m.group(2).strip()
        # Map: # CHƯƠNG / top title -> heading 1; ## -> 2 ; ### -> 3
        if txt.upper().startswith(("BÁO CÁO", "CHƯƠNG", "PHỤ LỤC", "LỜI NÓI", "TÀI LIỆU", "DANH MỤC")):
            hlevel = 1
        else:
            hlevel = min(level, 3)
        h = doc.add_heading(level=hlevel)
        add_inline(h, txt)
        i += 1
        continue

    # Blockquote
    if stripped.startswith(">"):
        txt = stripped.lstrip(">").strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        r_marker = p.add_run("│ ")
        add_inline(p, txt)
        for rr in p.runs:
            rr.italic = True
            rr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        i += 1
        continue

    # Bullet list
    if re.match(r"^[-*]\s+", stripped):
        txt = re.sub(r"^[-*]\s+", "", stripped)
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, txt)
        i += 1
        continue

    # Numbered list
    if re.match(r"^\d+\.\s+", stripped):
        txt = re.sub(r"^\d+\.\s+", "", stripped)
        p = doc.add_paragraph(style="List Number")
        add_inline(p, txt)
        i += 1
        continue

    # Empty line
    if stripped == "":
        i += 1
        continue

    # Image placeholder line -> highlighted centered text
    if stripped.startswith("📷"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(stripped)
        r.italic = True
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        i += 1
        continue

    # Normal paragraph
    p = doc.add_paragraph()
    add_inline(p, stripped)
    i += 1

doc.save(OUT)
print("SAVED", OUT)
